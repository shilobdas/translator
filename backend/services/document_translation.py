from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from .internal_config import internal_translate_text
from .translation import TranslationServiceError


DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
TXT_MIME_TYPE = "text/plain"


@dataclass
class TranslatedDocument:
    filename: str
    content: bytes
    mime_type: str
    target_language: str
    character_count: int


def translate_document_bytes(
    filename: str,
    content: bytes,
    source_language: str,
    target_languages: list[str],
) -> list[TranslatedDocument]:
    suffix = Path(filename).suffix.lower()
    if suffix == ".docx":
        return [
            _translate_docx(filename, content, source_language, target_language)
            for target_language in target_languages
        ]
    if suffix == ".txt":
        return [
            _translate_txt(filename, content, source_language, target_language)
            for target_language in target_languages
        ]
    raise TranslationServiceError("Unsupported document type. Upload .docx or .txt.")


def _translate_txt(
    filename: str,
    content: bytes,
    source_language: str,
    target_language: str,
) -> TranslatedDocument:
    text = _decode_text(content)
    translated_lines = []
    character_count = 0
    for line in text.splitlines(keepends=True):
        line_body = line.rstrip("\r\n")
        newline = line[len(line_body):]
        if not line_body.strip():
            translated_lines.append(line)
            continue
        result = internal_translate_text(line_body, source_language, target_language)
        translated_lines.append(result.text + newline)
        character_count += len(line_body)

    translated_text = "".join(translated_lines)
    return TranslatedDocument(
        filename=_translated_filename(filename, target_language),
        content=translated_text.encode("utf-8"),
        mime_type=TXT_MIME_TYPE,
        target_language=target_language,
        character_count=character_count,
    )


def _translate_docx(
    filename: str,
    content: bytes,
    source_language: str,
    target_language: str,
) -> TranslatedDocument:
    try:
        from docx import Document
    except ImportError as exc:
        raise TranslationServiceError(
            "python-docx is required for .docx translation. Install requirements.txt."
        ) from exc

    document = Document(BytesIO(content))
    character_count = 0
    for paragraph in _iter_paragraphs(document):
        text = paragraph.text
        if not text.strip():
            continue
        result = internal_translate_text(text, source_language, target_language)
        _replace_paragraph_text(paragraph, result.text)
        character_count += len(text)

    output = BytesIO()
    document.save(output)
    return TranslatedDocument(
        filename=_translated_filename(filename, target_language),
        content=output.getvalue(),
        mime_type=DOCX_MIME_TYPE,
        target_language=target_language,
        character_count=character_count,
    )


def _iter_paragraphs(document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def _replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise TranslationServiceError("Text file must be UTF-8 or Windows-1252 encoded.")


def _translated_filename(filename: str, target_language: str) -> str:
    path = Path(filename)
    return f"{path.stem}_{target_language}{path.suffix}"
